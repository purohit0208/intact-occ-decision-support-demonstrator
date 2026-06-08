from __future__ import annotations

import json
import shutil
import unicodedata
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from jatm_eval_core import (
    MANUSCRIPT_DIR,
    OUTPUT_FIGURES,
    OUTPUT_TABLES,
    REVISION_ROOT,
    REPORTS_DIR,
    REPO_ROOT,
    add_page_numbers,
    read_csv_safe,
    set_document_defaults,
    set_run_font,
)


TITLE_SELECTED = (
    "Pre-Arrival OCC Decision Support for Aircraft Turnaround Coordination: "
    "A Human-in-the-Loop Architecture and Scenario-Based Evaluation"
)

TITLE_OPTIONS = [
    TITLE_SELECTED,
    "Integrated OCC Decision Support for Aircraft Turnaround Operations Using Predictive Maintenance, Inventory Readiness, and Bottleneck Forecasting",
    "Human-in-the-Loop Pre-Arrival Decision Support for Aircraft Turnaround Operations: Architecture, Scenario Evaluation, and Evidence Boundaries",
]

FUNDING_TEXT = (
    "Parts of the research work presented in this paper were funded by the German Federal Ministry "
    "for Economic Affairs and Energy under grant 20D1931B (project CANARIA) as part of the LuFo VI-1 "
    "program and grant 20D2128E (project INTACT) as part of the LuFo VI-2 program."
)

ACKNOWLEDGMENTS_TEXT = (
    "The authors thank the German Federal Ministry for Economic Affairs and Energy for support under "
    "grant 20D1931B (project CANARIA) as part of the LuFo VI-1 program and grant 20D2128E "
    "(project INTACT) as part of the LuFo VI-2 program, as well as all partners participating in the "
    "CANARIA and INTACT research projects."
)


def clean_compact_df(df: pd.DataFrame) -> pd.DataFrame:
    compact = df.copy()
    compact = compact.replace({float("nan"): "N/A"})
    compact = compact.fillna("N/A")
    return compact.replace({"nan": "N/A", "NaN": "N/A", "": "N/A"})


def parse_holdout_metrics(raw: str) -> dict[str, float]:
    try:
        parsed = json.loads(str(raw))
    except Exception:
        return {}
    if isinstance(parsed, dict):
        return parsed
    return {}


def validation_metric(task: str, raw_metrics: str) -> str:
    metrics = parse_holdout_metrics(raw_metrics)
    if not metrics:
        return "N/A"
    if "roc_auc" in metrics:
        return f"ROC-AUC {float(metrics['roc_auc']):.3f}"
    if "f1_macro" in metrics and "roc_auc_ovr" in metrics:
        return f"macro-F1 {float(metrics['f1_macro']):.3f}; ROC-AUC OVR {float(metrics['roc_auc_ovr']):.3f}"
    if "f1_macro" in metrics:
        return f"macro-F1 {float(metrics['f1_macro']):.3f}"
    if "rmse" in metrics and "r2" in metrics:
        return f"RMSE {float(metrics['rmse']):.3f}; R2 {float(metrics['r2']):.3f}"
    return "; ".join(f"{key} {float(value):.3f}" for key, value in metrics.items())


def output_granularity(task: str) -> str:
    if "maintenance_failure" in task:
        return "Component-aware features; flight-scenario risk output"
    if "maintenance_urgency" in task:
        return "Component-aware features; flight-scenario urgency label"
    if "maintenance_remaining" in task:
        return "Flight-scenario remaining-flights proxy"
    if "inventory" in task:
        return "Service-area, trolley-level, and flight-level inventory-risk support"
    if "turnaround_bottleneck" in task:
        return "Process-class bottleneck probabilities"
    return "Flight-scenario advisory output"


def evidence_boundary(task: str) -> str:
    if "maintenance" in task:
        return "Synthetic-data support only; not certified maintenance planning"
    if "inventory" in task:
        return "Synthetic service-readiness risk; not item-wise demand or airline validation"
    if "bottleneck" in task:
        return "Synthetic process-class probabilities; not field duration validation"
    return "Synthetic scenario behavior only"


def display_task(task: str) -> str:
    labels = {
        "maintenance_failure_probability": "Maintenance failure risk",
        "maintenance_urgency_classification": "Maintenance urgency class",
        "maintenance_remaining_flights_regression": "Remaining-flights proxy",
        "inventory_risk_level_classification": "Inventory risk level",
        "turnaround_bottleneck_classification": "Bottleneck class",
    }
    return labels.get(task, task.replace("_", " "))


def display_regime(regime: str) -> str:
    labels = {
        "low_risk_routine": "Routine",
        "maintenance_heavy": "Maintenance",
        "inventory_heavy": "Inventory",
        "congestion_delay_heavy": "Congestion/delay",
        "assistance_readiness": "Assistance",
        "composite_multi_risk": "Composite",
    }
    return labels.get(regime, regime.replace("_", " "))


def rename_columns(df: pd.DataFrame, labels: dict[str, str]) -> pd.DataFrame:
    return df.rename(columns={col: labels.get(col, col.replace("_", " ")) for col in df.columns})


def marker_path(path: Path) -> str:
    try:
        return path.relative_to(REVISION_ROOT).as_posix()
    except ValueError:
        return path.as_posix()


def resolve_marker_path(path_text: str) -> Path:
    path = Path(path_text)
    if path.is_absolute():
        return path
    return REVISION_ROOT / path


def copy_and_make_figures() -> None:
    gui_source = REPO_ROOT / "paper_draft" / "jatm_submission" / "figures" / "Figure_5_dashboard_view.png"
    if gui_source.exists():
        shutil.copy2(gui_source, OUTPUT_FIGURES / "figure_6_representative_gui_dashboard.png")

    fig, ax = plt.subplots(figsize=(9.0, 4.8))
    ax.axis("off")
    columns = [
        ("Pre-arrival inputs", ["Cabin reports", "Inventory state", "Assistance request", "Delay and gate context"]),
        ("Analytical modules", ["Maintenance risk", "Inventory readiness", "Bottleneck class", "Assistance rules"]),
        ("OCC fusion", ["Readiness score", "Alert level", "Dominant bottleneck", "Action queue"]),
        ("Operator review", ["Coordinate teams", "Confirm resources", "Reassess after arrival", "Keep human control"]),
    ]
    x_positions = [0.04, 0.29, 0.54, 0.79]
    for x, (heading, items) in zip(x_positions, columns):
        outer = plt.Rectangle((x, 0.2), 0.18, 0.58, facecolor="#f7fafc", edgecolor="#2d3748", linewidth=1.2)
        ax.add_patch(outer)
        ax.text(x + 0.09, 0.72, heading, ha="center", va="center", fontsize=10, fontweight="bold")
        for idx, item in enumerate(items):
            y = 0.61 - idx * 0.10
            ax.text(x + 0.09, y, item, ha="center", va="center", fontsize=9)
    for start_x, end_x in zip(x_positions[:-1], x_positions[1:]):
        ax.annotate("", xy=(end_x, 0.49), xytext=(start_x + 0.18, 0.49), arrowprops={"arrowstyle": "->", "lw": 1.4, "color": "#2d3748"})
    ax.text(
        0.5,
        0.08,
        "Controlled synthetic scenarios evaluate architecture behavior; no airline deployment or operational validation is claimed.",
        ha="center",
        fontsize=9,
    )
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "figure_1_occ_architecture.png", dpi=300)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    ax.axis("off")
    boxes = [
        (0.03, 0.62, 0.19, 0.22, "Pre-arrival\nscenario state"),
        (0.29, 0.74, 0.18, 0.16, "Maintenance\nrisk module"),
        (0.29, 0.50, 0.18, 0.16, "Inventory\nreadiness module"),
        (0.29, 0.26, 0.18, 0.16, "Bottleneck\nforecast module"),
        (0.55, 0.62, 0.18, 0.22, "Transparent\nOCC fusion"),
        (0.79, 0.62, 0.18, 0.22, "Readiness,\nalerts, actions"),
    ]
    for x, y, w, h, text in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor="#edf2f7", edgecolor="#2d3748", linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10)
    arrows = [
        ((0.22, 0.73), (0.29, 0.82)),
        ((0.22, 0.73), (0.29, 0.58)),
        ((0.22, 0.73), (0.29, 0.34)),
        ((0.47, 0.82), (0.55, 0.73)),
        ((0.47, 0.58), (0.55, 0.73)),
        ((0.47, 0.34), (0.55, 0.73)),
        ((0.73, 0.73), (0.79, 0.73)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops={"arrowstyle": "->", "lw": 1.2, "color": "#2d3748"})
    ax.text(0.29, 0.10, "Passenger assistance is rule-based readiness support, not localization ML.", fontsize=9)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "figure_2_module_to_fusion_workflow.png", dpi=300)
    plt.close(fig)


def create_compact_tables() -> dict[str, Path]:
    tables: dict[str, Path] = {}
    module_source = read_csv_safe(REPO_ROOT / "paper_assets" / "tables" / "module_evaluation_summary.csv")
    if not module_source.empty:
        compact = pd.DataFrame(
            {
                "Module": module_source["module"],
                "Task": [display_task(task) for task in module_source["task"]],
                "Selected model": module_source["selected_model"],
                "Validation metric": [
                    validation_metric(task, raw)
                    for task, raw in zip(module_source["task"], module_source["holdout_metrics"])
                ],
                "Output granularity": [output_granularity(task) for task in module_source["task"]],
                "Evidence boundary": [evidence_boundary(task) for task in module_source["task"]],
            }
        )
        compact = clean_compact_df(compact)
        target = OUTPUT_TABLES / "manuscript_table_1_module_summary.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["module"] = target

    sweep_summary = read_csv_safe(OUTPUT_TABLES / "scenario_sweep_summary.csv")
    if not sweep_summary.empty:
        wanted = [
            "scenario_count",
            "mean_readiness",
            "median_readiness",
            "alert_routine_count",
            "alert_watch_count",
            "alert_priority_count",
            "alert_critical_count",
            "dominant_bottleneck_maintenance_count",
            "dominant_bottleneck_catering_count",
            "dominant_bottleneck_boarding_count",
            "mean_action_queue_size",
        ]
        compact = sweep_summary[sweep_summary["metric"].isin(wanted)].copy()
        compact["metric"] = compact["metric"].str.replace("_", " ", regex=False)
        compact = clean_compact_df(rename_columns(compact, {"metric": "Metric", "value": "Value"}))
        target = OUTPUT_TABLES / "manuscript_table_2_scenario_sweep_summary.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["sweep"] = target

    regime = read_csv_safe(OUTPUT_TABLES / "regime_summary_table.csv")
    if not regime.empty:
        columns = [
            "regime",
            "scenario_count",
            "mean_readiness",
            "median_readiness",
            "dominant_bottleneck_mode",
            "mean_maintenance_failure_probability",
            "mean_inventory_shortage_risk",
            "assistance_confirmed_minus_unconfirmed",
        ]
        compact = regime[[col for col in columns if col in regime.columns]].copy()
        compact["regime"] = compact["regime"].map(display_regime)
        compact = clean_compact_df(
            rename_columns(
                compact,
                {
                    "regime": "Regime",
                    "scenario_count": "n",
                    "mean_readiness": "Mean readiness",
                    "median_readiness": "Median readiness",
                    "dominant_bottleneck_mode": "Dominant bottleneck",
                    "mean_maintenance_failure_probability": "Maint. risk",
                    "mean_inventory_shortage_risk": "Inv. risk",
                    "assistance_confirmed_minus_unconfirmed": "Assist. effect",
                },
            )
        )
        target = OUTPUT_TABLES / "manuscript_table_3_regime_summary.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["regime"] = target

    monotonic = read_csv_safe(OUTPUT_TABLES / "monotonicity_summary.csv")
    if not monotonic.empty:
        compact = monotonic[["test", "passed", "criterion"]].copy()
        compact["test"] = compact["test"].str.replace("_", " ", regex=False)
        compact = clean_compact_df(rename_columns(compact, {"test": "Test", "passed": "Passed", "criterion": "Criterion"}))
        target = OUTPUT_TABLES / "manuscript_table_4_monotonicity_summary.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["monotonicity"] = target

    ablation = read_csv_safe(OUTPUT_TABLES / "ablation_summary.csv")
    if not ablation.empty:
        columns = [
            "configuration",
            "mean_readiness",
            "mean_readiness_delta_from_full",
            "critical_cases",
            "priority_cases",
            "mean_action_queue_size",
            "top_action_changed_cases",
        ]
        compact = ablation[columns].copy()
        compact["configuration"] = compact["configuration"].str.replace("_", " ", regex=False)
        compact = clean_compact_df(
            rename_columns(
                compact,
                {
                    "configuration": "Configuration",
                    "mean_readiness": "Mean readiness",
                    "mean_readiness_delta_from_full": "Delta from full",
                    "critical_cases": "Critical",
                    "priority_cases": "Priority",
                    "mean_action_queue_size": "Mean actions",
                    "top_action_changed_cases": "Top action changes",
                },
            )
        )
        target = OUTPUT_TABLES / "manuscript_table_5_ablation_summary.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["ablation"] = target

    baseline = read_csv_safe(OUTPUT_TABLES / "baseline_comparison.csv")
    if not baseline.empty:
        columns = [
            "baseline",
            "mean_readiness",
            "std_readiness",
            "unique_alert_levels",
            "unique_dominant_signals",
            "unique_action_categories",
            "critical_cases",
            "priority_cases",
        ]
        compact = baseline[columns].copy()
        compact["baseline"] = compact["baseline"].str.replace("_", " ", regex=False)
        compact = clean_compact_df(
            rename_columns(
                compact,
                {
                    "baseline": "Baseline",
                    "mean_readiness": "Mean readiness",
                    "std_readiness": "Std readiness",
                    "unique_alert_levels": "Alert levels",
                    "unique_dominant_signals": "Dominant signals",
                    "unique_action_categories": "Action categories",
                    "critical_cases": "Critical",
                    "priority_cases": "Priority",
                },
            )
        )
        target = OUTPUT_TABLES / "manuscript_table_6_baseline_comparison.csv"
        compact.to_csv(target, index=False, na_rep="N/A")
        tables["baseline"] = target
    return tables


def metric(metric_name: str) -> str:
    summary = read_csv_safe(OUTPUT_TABLES / "scenario_sweep_summary.csv")
    if summary.empty:
        return "NA"
    hit = summary[summary["metric"] == metric_name]
    if len(hit):
        value = hit["value"].iat[0]
        try:
            numeric = float(value)
            if numeric.is_integer():
                return str(int(numeric))
            return f"{numeric:.3f}".rstrip("0").rstrip(".")
        except Exception:
            return str(value)
    return "NA"


def first_value(path: Path, key_col: str, key: str, value_col: str) -> str:
    df = read_csv_safe(path)
    if df.empty:
        return "NA"
    hit = df[df[key_col] == key]
    if len(hit):
        value = hit[value_col].iat[0]
        try:
            numeric = float(value)
            return f"{numeric:.2f}".rstrip("0").rstrip(".")
        except Exception:
            return str(value)
    return "NA"


def clean_references(text: str) -> str:
    try:
        text = text.encode("latin1").decode("utf-8")
    except UnicodeError:
        pass
    text = unicodedata.normalize("NFKD", text)
    return text.encode("ascii", "ignore").decode("ascii").strip()


def reference_block() -> str:
    local = MANUSCRIPT_DIR / "JATM_References.md"
    if local.exists():
        return clean_references(local.read_text(encoding="utf-8", errors="replace"))
    source = REPO_ROOT / "paper_draft" / "jatm_submission" / "jatm_main_manuscript.md"
    if source.exists():
        text = source.read_text(encoding="utf-8", errors="replace")
        if "## References" in text:
            references = clean_references(text.split("## References", 1)[1])
            local.write_text(references + "\n", encoding="utf-8")
            return references
    return "References should be completed from the manuscript source bibliography before submission."


def build_markdown(tables: dict[str, Path]) -> Path:
    scenario_count = metric("scenario_count")
    mean_readiness = metric("mean_readiness")
    median_readiness = metric("median_readiness")
    critical_count = metric("alert_critical_count")
    priority_count = metric("alert_priority_count")
    routine_count = metric("alert_routine_count")
    composite_mean = first_value(OUTPUT_TABLES / "regime_summary_table.csv", "regime", "composite_multi_risk", "mean_readiness")
    low_mean = first_value(OUTPUT_TABLES / "regime_summary_table.csv", "regime", "low_risk_routine", "mean_readiness")
    maintenance_mean = first_value(OUTPUT_TABLES / "regime_summary_table.csv", "regime", "maintenance_heavy", "mean_readiness")
    inventory_mean = first_value(OUTPUT_TABLES / "regime_summary_table.csv", "regime", "inventory_heavy", "mean_readiness")
    ablation_no_maintenance = first_value(OUTPUT_TABLES / "ablation_summary.csv", "configuration", "no_maintenance_in_fusion", "mean_readiness_delta_from_full")
    ablation_no_bottleneck = first_value(OUTPUT_TABLES / "ablation_summary.csv", "configuration", "no_bottleneck_in_fusion", "mean_readiness_delta_from_full")

    body = f"""# {TITLE_SELECTED}

## Abstract

Aircraft turnaround coordination depends on maintenance, catering, passenger-assistance, gate, and bottleneck information that may become available before arrival but is often handled through fragmented workflows. This paper presents a human-in-the-loop Operations Control Center (OCC) decision-support architecture for pre-arrival aircraft turnaround coordination. The architecture integrates component-aware maintenance risk indicators, service-area inventory readiness, process-class bottleneck forecasting, passenger-assistance readiness logic, and a transparent rule-based fusion layer that converts heterogeneous signals into readiness scores, alert levels, and operator-facing action queues. The evaluation is deliberately framed as controlled synthetic scenario analysis rather than airline deployment validation. Existing module-level holdout results are combined with a deterministic scenario sweep of {scenario_count} synthetic inbound cases across six operational regimes, monotonicity and sanity checks, and ablation and baseline comparisons. The scenario sweep produced mean readiness {mean_readiness}, median readiness {median_readiness}, {critical_count} Critical cases, {priority_count} Priority cases, and {routine_count} Routine cases. Ablations show that removing maintenance influence increased mean readiness by {ablation_no_maintenance} points, while removing bottleneck influence increased mean readiness by {ablation_no_bottleneck} points, indicating that the fusion layer materially changes prioritization under controlled stress conditions. Two maintenance-related monotonicity checks did not strictly pass because small local readiness increases occurred at low malfunction levels before urgency thresholds changed; these failures are treated as limitations of the current model-fusion interaction rather than hidden. The work contributes a reproducible OCC decision-support architecture and scenario-evaluation package, while explicitly preserving the limitations that no airline operational data, OCC user study, deployment, certification, or measured operational savings are claimed.

**Keywords:** aircraft turnaround; Operations Control Center; decision support; scenario-based evaluation; predictive maintenance; inventory readiness; bottleneck forecasting; human-in-the-loop

## 1. Introduction

Aircraft turnaround operations are managed under tight time constraints and require coordination across maintenance, catering, cleaning, boarding, refueling, passenger assistance, gate management, and airline operations control. Information relevant to these processes can be available before arrival, but it is often distributed across separate operational or analytical views. The practical problem is therefore not only whether individual prediction models can be trained, but how heterogeneous pre-arrival signals can be converted into a form that supports OCC coordination before the aircraft reaches the stand.

This study is framed as an air transport operations contribution. The focus is pre-arrival decision support for turnaround coordination, not a generic intelligent-systems contribution. The architecture remains human-in-the-loop: it supports operator review through readiness scores, alert levels, provenance labels, and action queues, but it does not automate dispatch or claim certified operational control.

The work uses synthetic but operationally grounded data because no real airline, airport, or station-operation dataset is available in the current project package. That boundary is central to the paper. The evaluation therefore asks whether the architecture behaves coherently and differentiates operational regimes under controlled scenarios. It does not ask whether the system is validated for airline deployment.

## 2. Related Work

### 2.1 OCC Decision Support and Disruption Management

Airline Operations Control Centers coordinate time-critical decisions across aircraft, crews, gates, passengers, and downstream rotations. Classical OCC and disruption-management literature, including work by Clarke, Clausen, Hassan, and Su, shows that these decisions combine schedule recovery, resource coordination, and human judgment rather than a single optimization objective. Naturalistic and decision-support studies by Fogaca and Vink further motivate systems that preserve operator review instead of replacing OCC staff with automated control.

### 2.2 Turnaround Coordination and Ground Handling

Aircraft turnaround research treats the stand-side process as a tightly coupled operational system. Work on turnaround management, ground-handling automation, apron-resource planning, and delay propagation, including studies by Makhloof, Conde, Gok, Guimarans, Evler, and Wu, shows that local process delays can affect broader airline and airport performance. This paper builds on that perspective by focusing on pre-arrival coordination signals that can change preparation before the aircraft reaches the stand.

### 2.3 Aviation ML for Delay and Resource Prediction

Aviation machine-learning studies have addressed arrival-time prediction, delay classification, resource planning, and ground-operation forecasting. The cited studies by Basturk, Zoutendijk, Sahadevan, Okwir, Kicinger, and Wandelt support the view that predictive analytics can inform aviation operations. However, a collection of separate predictive tasks does not automatically create an OCC workflow. The contribution here is therefore not a new prediction algorithm; it is the operational integration of heterogeneous predictive and rule-based signals into one pre-arrival advisory process.

### 2.4 Predictive Maintenance, Service Readiness, and Assistance

Predictive-maintenance studies by Dangut, Lee and Mitici, Zeng, Hakami, and Kabashkin motivate maintenance-risk indicators for aviation decision support. Adjacent work on food-service and demand forecasting, including Malefors and Rodrigues, is relevant to service-readiness reasoning, but it should not be interpreted as evidence for item-wise airline trolley-demand prediction in this paper. Passenger-assistance and accessibility work by Gotti and Lee motivates explicit assistance-readiness handling, while the implemented system treats assistance as rule-based readiness support rather than localization machine learning.

### 2.5 Human-in-the-Loop and Synthetic Scenario Evaluation

Human-in-the-loop AI and explainability work, including Mosqueira-Rey, Vaccaro, and Xie, supports transparent advisory systems where model outputs remain reviewable by human operators. Synthetic-data and scenario-evaluation studies by Ardebili, de Frutos, and Stefani provide a methodological basis for controlled evaluation when operational data are not available, but they do not remove the need for future airline data and user studies. This paper therefore uses synthetic scenarios to evaluate architecture behavior, not to claim field validity.

Across these streams, the open gap is the connection between module-level analytics and pre-arrival OCC coordination. Maintenance, inventory, bottleneck, and assistance-readiness signals are not treated as separate dashboards. They are fused into one operator-facing advisory workflow, with explicit labels distinguishing trained inference, rule-based readiness logic, and synthetic scenario simulation.

## 3. Contributions

This paper makes four bounded contributions:

- A pre-arrival OCC decision-support architecture for aircraft turnaround coordination.
- A modular integration of maintenance risk, inventory readiness, bottleneck forecasting, and passenger-assistance readiness signals.
- A transparent fusion/action layer that converts heterogeneous signals into readiness scores, alerts, and action queues for operator review.
- A scenario-based evaluation package with seeded cases, a {scenario_count}-case synthetic sweep, monotonic sanity checks, and ablation/baseline comparisons.

The contribution is architectural and operational. It is not a claim of a new AI algorithm, airline deployment, or field-validated performance.

## 4. System Architecture

The architecture has four layers. The input layer represents pre-arrival aircraft and turnaround state, including passenger load, arrival delay, gate congestion, malfunction count and severity, trolley availability, inventory shortage indicators, assistance requests, weather disturbance, and refueling requirement. The analytics layer contains trained local modules for maintenance, inventory, and bottleneck prediction. Passenger assistance is treated as rule-based readiness support, not localization machine learning. The fusion layer combines module outputs and scenario state into readiness, criticality, alerts, and action queues. The presentation layer exposes the result in an operator-facing OCC interface.

[[FIGURE:{marker_path(OUTPUT_FIGURES / "figure_1_occ_architecture.png")}|Figure 1. Integrated OCC decision-support architecture for pre-arrival turnaround coordination.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "figure_2_module_to_fusion_workflow.png")}|Figure 2. Module-to-fusion workflow for pre-arrival OCC decision support.]]

## 5. Module Design and Evidence Boundary

The maintenance module is component-aware in its features and explanation logic, but the OCC interface summarizes the result at flight-scenario level. It should not be described as a certified component-wise aircraft-health dashboard. The inventory module supports service-area, trolley-level, and flight-level inventory-risk reasoning; it is not item-wise product demand prediction. The bottleneck module predicts process-class bottleneck probabilities across Maintenance, Catering, Cleaning, Boarding, and Refueling. Passenger assistance is rule-based readiness support.

[[TABLE:{marker_path(tables.get("module", Path()))}|Table 1. Module-level model summary and evidence boundary.]]

## 6. Scenario-Sweep Design

The evaluation extends the five seeded demonstration cases with a deterministic scenario sweep. The sweep contains {scenario_count} synthetic inbound cases generated under six regimes: low-risk routine, maintenance-heavy, inventory-heavy, congestion/delay-heavy, assistance-readiness, and composite multi-risk. A fixed random seed, 20260608, is used so that the inputs, outputs, and figures can be reproduced. Cases are assigned across the six regimes and then evaluated through the existing trained model inference and OCC fusion functions.

The generator varies arrival delay, gate congestion, malfunction count, malfunction severity, trolley availability, inventory shortage, passenger load, assistance readiness, weather disturbance, and refueling requirement. Low-risk routine cases use low delay, low congestion, few or no malfunctions, adequate trolley availability, and confirmed or absent assistance requests. Maintenance-heavy cases stress malfunction count and severity. Inventory-heavy cases stress trolley availability and shortage indicators. Congestion/delay-heavy cases stress arrival delay and gate congestion. Assistance-readiness cases vary assistance request status and confirmation. Composite multi-risk cases combine multiple stressors.

For each generated inbound case, the evaluation records readiness score, alert level, dominant bottleneck, module-level risks, action-queue size, priority-action counts, and top action categories. The sweep is a controlled architecture-behavior evaluation. It does not replace the seeded demonstration cases and does not create observational airline evidence.

[[TABLE:{marker_path(tables.get("sweep", Path()))}|Table 2. Overall scenario-sweep summary.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "scenario_sweep_readiness_distribution.png")}|Figure 3. Readiness distribution across the controlled synthetic scenario sweep.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "scenario_sweep_alert_distribution.png")}|Figure 4. Alert-level distribution across the controlled synthetic scenario sweep.]]

## 7. Regime-Level Results

The generated regimes produced differentiated system behavior. The low-risk routine regime had mean readiness {low_mean}, while the maintenance-heavy regime had mean readiness {maintenance_mean}, the inventory-heavy regime had mean readiness {inventory_mean}, and the composite multi-risk regime had mean readiness {composite_mean}. These values should be interpreted as designed scenario-response behavior, not as observed airline probabilities.

[[TABLE:{marker_path(tables.get("regime", Path()))}|Table 3. Regime-level summary of synthetic scenario responses.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "regime_readiness_boxplot.png")}|Figure 5. Readiness by synthetic operational regime.]]

## 8. Monotonicity and Sanity Checks

The monotonic checks test whether the system behaves coherently when one stressor changes while other inputs are held constant. Six of eight checks passed. Inventory shortage, trolley availability, arrival delay, gate congestion, assistance readiness, and composite-versus-low-risk checks behaved in the expected direction. The malfunction-severity and malfunction-count checks did not strictly pass because readiness increased slightly at low malfunction values before the maintenance urgency threshold changed. In both cases, maintenance failure probability still increased with the stressor, but the fused readiness score was not perfectly monotonic.

These failures are not positive evidence. They show that the current model-fusion interaction is not strictly monotonic under every stressor sweep and would need additional monotonic post-model guards, calibrated fusion constraints, or expert-reviewed thresholds before any deployment-oriented study. The present paper reports the failure as a limitation rather than converting the system into unsupported field evidence.

[[TABLE:{marker_path(tables.get("monotonicity", Path()))}|Table 4. Monotonicity and sanity-check summary.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "sensitivity_curves.png")}|Figure 6. Sensitivity curves for one-factor system-behavior checks.]]

## 9. Ablation and Baseline Results

The ablation study compares the full fusion system against configurations in which maintenance, inventory, bottleneck, or assistance influence is neutralized. The objective is not to prove real operational superiority. The objective is to show whether each signal family changes the behavior of the integrated action logic under controlled synthetic scenarios.

The full configuration uses maintenance, inventory, bottleneck, assistance, and fusion logic together. The no-maintenance ablation neutralizes maintenance influence in the fusion layer. The no-inventory ablation neutralizes inventory-shortage influence. The no-bottleneck ablation replaces process-class bottleneck evidence with a neutral bottleneck state. The no-assistance ablation removes assistance-readiness influence. Each configuration is evaluated on the same generated scenario inputs so differences reflect fusion behavior rather than different scenario samples.

[[TABLE:{marker_path(tables.get("ablation", Path()))}|Table 5. Ablation summary relative to the full integrated system.]]

[[FIGURE:{marker_path(OUTPUT_FIGURES / "ablation_readiness_delta.png")}|Figure 7. Mean readiness change when signal families are removed from fusion.]]

The baseline comparison contrasts the full system with static delay/gate, maintenance-only, inventory-only, and bottleneck-only baselines. The static delay/gate baseline uses only arrival delay and gate congestion. The single-signal baselines expose one module family at a time. The full system produces more cross-domain action differentiation than the single-signal baselines, but this should be read as scenario-behavior differentiation rather than field performance superiority.

[[TABLE:{marker_path(tables.get("baseline", Path()))}|Table 6. Baseline comparison for controlled synthetic behavior differentiation.]]

## 10. Operator Interface

The interface remains a supporting artifact, not the main evidence source. It shows how readiness, alerts, provenance, bottleneck state, and action queues can be presented for operator review. Only one representative screenshot is retained in the main paper to avoid making the manuscript look like a software demo report.

[[FIGURE:{marker_path(OUTPUT_FIGURES / "figure_6_representative_gui_dashboard.png")}|Figure 8. Representative OCC dashboard view showing fused pre-arrival summaries.]]

## 11. Discussion

The results support the architectural claim that heterogeneous pre-arrival signals can be fused into differentiated OCC advisory outputs. Maintenance-heavy, inventory-heavy, congestion-heavy, assistance-readiness, routine, and composite scenarios lead to distinct readiness states, alert distributions, bottleneck profiles, and action categories. This is the appropriate contribution level for the current evidence base.

For air transport management readers, the practical implication is not that the demonstrator is ready for airline use. The implication is that pre-arrival information can be structured into a transparent workflow that helps reason about cross-domain turnaround readiness before arrival. A maintenance issue can change stand-side preparation, an inventory issue can affect catering and boarding readiness, and an unconfirmed assistance request can reduce coordination margin even when other modules are nominal.

## 12. Limitations and Threats to Validity

The main limitation is synthetic data. The scenario sweep is controlled and reproducible, but it is not observational airline evidence. The system has not been evaluated on real airline, airport, or station-operation data. No OCC user study has been conducted. No airline deployment, certification, safety assurance, operational savings, or return-on-investment claim is made.

The readiness thresholds are heuristic. The fusion layer is transparent but rule-based rather than learned from OCC decisions. The maintenance and inventory modules are trained on synthetic labels. The bottleneck model predicts process classes, not full turnaround duration. Passenger assistance is readiness logic, not localization ML. The failed maintenance monotonicity checks further show that the current trained-module plus fusion behavior would need additional hardening before any practical deployment study.

## 13. Conclusion

This paper presents a pre-arrival OCC decision-support architecture for aircraft turnaround coordination and evaluates it through controlled synthetic scenarios. The results show that heterogeneous pre-arrival maintenance, inventory, bottleneck, and assistance-readiness signals can be fused into differentiated advisory outputs while preserving a transparent human-in-the-loop workflow. The study also shows the limits of the present evidence base: controlled scenario behavior is useful for architecture evaluation, but it is not a substitute for airline data, expert review, OCC user evaluation, or deployment-oriented validation.

## Data, Code, and Artifact Availability

The demonstrator source code, scenario-evaluation scripts, generated tables, and generated figures used in this study are available at https://github.com/purohit0208/intact-occ-decision-support-demonstrator. No airline operational dataset is included.

## Declaration of Generative AI and AI-Assisted Technologies

During manuscript preparation, OpenAI Codex was used to support document structuring, local script generation, and formatting. The authors are responsible for reviewing, verifying, and approving the final manuscript content before submission.

## References

{reference_block()}
"""
    path = MANUSCRIPT_DIR / "JATM_Main_Manuscript.md"
    path.write_text(body, encoding="utf-8")
    return path


def sanitize(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_csv_table(document: Document, csv_path: Path, caption: str) -> None:
    table_df = pd.read_csv(csv_path, keep_default_na=False).astype(str)
    rows = list(table_df.itertuples(index=False, name=None))
    headers = list(table_df.columns)
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        cell.text = sanitize(str(header))
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=7, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = sanitize(str(value))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=7)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cap.add_run(caption), size=9, italic=True)


def add_figure(document: Document, fig_path: Path, caption: str) -> None:
    if not fig_path.exists():
        document.add_paragraph(f"[Missing figure: {fig_path}]")
        return
    document.add_picture(str(fig_path), width=Inches(5.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cap.add_run(caption), size=9, italic=True)


def flush_paragraph(document: Document, buffer: list[str], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    if not buffer:
        return
    text = sanitize(" ".join(part.strip() for part in buffer if part.strip()))
    if text:
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        set_run_font(paragraph.add_run(text), size=10)
    buffer.clear()


def build_main_docx(markdown_path: Path) -> Path:
    document = Document()
    set_document_defaults(document)
    add_page_numbers(document)
    buffer: list[str] = []
    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            flush_paragraph(document, buffer, alignment)
            continue
        if stripped.startswith("[[TABLE:") and stripped.endswith("]]"):
            flush_paragraph(document, buffer, alignment)
            payload = stripped[len("[[TABLE:") : -2]
            path_text, caption = payload.split("|", 1)
            add_csv_table(document, resolve_marker_path(path_text), caption)
            continue
        if stripped.startswith("[[FIGURE:") and stripped.endswith("]]"):
            flush_paragraph(document, buffer, alignment)
            payload = stripped[len("[[FIGURE:") : -2]
            path_text, caption = payload.split("|", 1)
            add_figure(document, resolve_marker_path(path_text), caption)
            continue
        if stripped.startswith("# "):
            flush_paragraph(document, buffer, alignment)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(sanitize(stripped[2:])), size=14, bold=True)
            continue
        if stripped.startswith("### "):
            flush_paragraph(document, buffer, alignment)
            document.add_heading(sanitize(stripped[4:].strip()), level=2)
            continue
        if stripped.startswith("## "):
            flush_paragraph(document, buffer, alignment)
            heading = stripped[3:].strip()
            document.add_heading(sanitize(heading), level=1)
            alignment = WD_ALIGN_PARAGRAPH.LEFT if heading == "References" else WD_ALIGN_PARAGRAPH.JUSTIFY
            continue
        if stripped.startswith("- "):
            flush_paragraph(document, buffer, alignment)
            paragraph = document.add_paragraph(style="List Bullet")
            set_run_font(paragraph.add_run(sanitize(stripped[2:])), size=10)
            continue
        buffer.append(stripped)
    flush_paragraph(document, buffer, alignment)
    target = MANUSCRIPT_DIR / "JATM_Main_Manuscript_Anonymized.docx"
    document.save(target)
    return target


def build_title_page() -> Path:
    document = Document()
    set_document_defaults(document)
    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run(TITLE_SELECTED), size=12, bold=True)
    sections = [
        ("Target journal", ["Journal of Air Transport Management"]),
        ("Author names", ["Parth Yogeshbhai Purohit", "Thomas Feuerle", "Peter Hecker"]),
        (
            "Affiliations",
            [
                "Technische Universität Braunschweig, Institut für Flugführung | Institute of Flight Guidance, Hermann-Blenk-Str. 27, 38108 Braunschweig, Germany"
            ],
        ),
        (
            "Corresponding author",
            [
                "Name: Parth Yogeshbhai Purohit",
                "Email: parth-yogeshbhai.purohit@tu-braunschweig.de",
                "ORCID: 0009-0005-1547-8992",
                "Address: Technische Universität Braunschweig, Institut für Flugführung | Institute of Flight Guidance, Hermann-Blenk-Str. 27, 38108 Braunschweig, Germany",
                "Website: https://www.tu-braunschweig.de/iff",
            ],
        ),
        ("Funding", [FUNDING_TEXT]),
        ("Acknowledgments", [ACKNOWLEDGMENTS_TEXT]),
        (
            "Declaration of competing interests",
            [
                "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."
            ],
        ),
        (
            "Author contributions",
            [
                "Parth Yogeshbhai Purohit: Conceptualization, Methodology, Software, Validation, Investigation, Writing - original draft, Visualization.",
                "Thomas Feuerle: Formal analysis, Writing - review & editing, Supervision, Project administration, Funding acquisition.",
                "Peter Hecker: Formal analysis, Writing - review & editing, Project administration, Funding acquisition.",
            ],
        ),
    ]
    for heading, lines in sections:
        p = document.add_paragraph()
        set_run_font(p.add_run(heading), size=9, bold=True)
        for line in lines:
            paragraph = document.add_paragraph()
            set_run_font(paragraph.add_run(line), size=9)
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
    target = MANUSCRIPT_DIR / "JATM_Title_Page.docx"
    document.save(target)
    return target


def build_highlights() -> Path:
    highlights = [
        "Pre-arrival operations control support is tested for aircraft turnaround.",
        "Maintenance, inventory, bottleneck, and assistance signals are fused.",
        "One thousand synthetic cases cover six operational regimes.",
        "Signal-removal tests show how inputs change readiness and actions.",
        "Limitations state no airline data, deployment, user study, or savings claim.",
    ]
    (MANUSCRIPT_DIR / "JATM_Highlights.md").write_text("\n".join(f"- {item}" for item in highlights) + "\n", encoding="utf-8")
    document = Document()
    set_document_defaults(document)
    p = document.add_paragraph()
    set_run_font(p.add_run("Highlights"), size=12, bold=True)
    for item in highlights:
        paragraph = document.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(item), size=10)
    target = MANUSCRIPT_DIR / "JATM_Highlights.docx"
    document.save(target)
    return target


def write_reports() -> None:
    monotonic = read_csv_safe(OUTPUT_TABLES / "monotonicity_summary.csv")
    failures = []
    if not monotonic.empty:
        failures = list(monotonic[monotonic["passed"] == False]["test"])  # noqa: E712

    title_options = "\n".join(f"{idx + 1}. {title}" for idx, title in enumerate(TITLE_OPTIONS))
    (MANUSCRIPT_DIR / "JATM_Title_Options.md").write_text(f"# JATM Title Options\n\n{title_options}\n", encoding="utf-8")

    technical = f"""# Technical Change Report

## What changed

- Created an isolated JATM revision workspace separate from the live demonstrator.
- Added deterministic evaluation scripts for a 1000-case scenario sweep, regime summaries, monotonicity checks, ablations, and baselines.
- Generated new CSV tables and manuscript-scale figures under the revision folder.
- Reframed the manuscript for Journal of Air Transport Management around pre-arrival OCC decision support and aircraft turnaround coordination.
- Reduced software-demo emphasis by retaining one representative GUI screenshot and adding scenario/evaluation figures.
- Removed reviewer-visible revision/meta language, replaced local-path availability text with conservative repository/on-request wording, expanded Related Work and methods detail, enriched Table 1, and normalized missing values as `N/A`.
- Preserved strong limitations around synthetic data, no OCC user study, no airline data, no deployment, no certification, and no measured operational savings.

## What did not change

- `Start-OCC-Demo.bat`, `Stop-OCC-Demo.bat`, backend services, frontend source, saved model bundles, and the five live seeded demo scenarios were not intentionally edited.
- The new evaluation layer imports existing service functions as read-only dependencies and writes outputs only inside `paper_revision_jatm`.

## Monotonicity audit

Passed checks: {0 if monotonic.empty else int(monotonic['passed'].sum())}/{0 if monotonic.empty else len(monotonic)}.
Failed checks reported in the manuscript: {', '.join(failures) if failures else 'none'}.
"""
    (REPORTS_DIR / "technical_change_report.md").write_text(technical, encoding="utf-8")

    revision_note = """# Presentation and Contribution Revision Note

The JATM revision strengthens presentation and contribution framing without changing the evidence boundary.

## Presentation

- The paper is reframed for air transport operations rather than generic intelligent systems.
- The figure set is simplified: architecture, module-to-fusion workflow, scenario distributions, ablation/sensitivity results, and one representative interface screenshot.
- Large screenshot-heavy presentation is reduced.

## Contribution

- The contribution is sharpened as a pre-arrival OCC decision-support architecture for turnaround coordination.
- The five live demo scenarios are preserved but no longer carry the full evaluation burden.
- A separate 1000-case synthetic scenario sweep, regime analysis, monotonicity checks, ablations, and baselines provide stronger controlled evidence.

## Evidence boundary

- The revision does not claim real-world validation, OCC user acceptance, airline deployment, certification, or operational savings.
- Synthetic-data limitations are explicit and central.
"""
    (REPORTS_DIR / "presentation_contribution_revision_note.md").write_text(revision_note, encoding="utf-8")

    recommendation = """# JATM Go/No-Go Recommendation

Recommendation: conditional go for JATM after author metadata is completed and the supervisor approves the conservative synthetic-evaluation framing.

Rationale:

- JATM is a strong scope fit because the paper is fundamentally about air transport operations, OCC workflow, and aircraft turnaround coordination.
- The scenario sweep and ablation package materially strengthen the paper compared with a five-scenario demonstrator-only evaluation.
- The paper remains at risk because it uses synthetic data and has no OCC user study or airline operational validation.

Do not submit if the intended claim is real operational performance. Submit only as a controlled, reproducible, scenario-based decision-support architecture study.
"""
    (REPORTS_DIR / "go_no_go_recommendation.md").write_text(recommendation, encoding="utf-8")

    package_readme = """# JATM Revision Manuscript Package

Files:

- `JATM_Main_Manuscript.md`
- `JATM_Main_Manuscript_Anonymized.docx`
- `JATM_Title_Page.docx`
- `JATM_Highlights.docx`
- `JATM_Title_Options.md`

Before submission:

- Confirm that the public GitHub repository includes the `paper_revision_jatm` artifact folder before submission.
- Recheck all author affiliations.
- Confirm the author declaration and AI-use declaration text.
- Confirm with the supervisor that synthetic scenario evaluation is acceptable for the chosen target.
"""
    (MANUSCRIPT_DIR / "README_JATM_revision_package.md").write_text(package_readme, encoding="utf-8")


def main() -> None:
    copy_and_make_figures()
    tables = create_compact_tables()
    markdown = build_markdown(tables)
    docx = build_main_docx(markdown)
    title = build_title_page()
    highlights = build_highlights()
    write_reports()
    print(f"Wrote {markdown}")
    print(f"Wrote {docx}")
    print(f"Wrote {title}")
    print(f"Wrote {highlights}")
    print(f"Wrote reports under {REPORTS_DIR}")


if __name__ == "__main__":
    main()
