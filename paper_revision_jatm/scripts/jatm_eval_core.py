from __future__ import annotations

import csv
import hashlib
import json
import math
import random
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from statistics import mean, median
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REVISION_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = REVISION_ROOT.parent
BACKEND_ROOT = REPO_ROOT / "backend"
OUTPUT_TABLES = REVISION_ROOT / "outputs" / "tables"
OUTPUT_FIGURES = REVISION_ROOT / "outputs" / "figures"
REPORTS_DIR = REVISION_ROOT / "reports"
MANUSCRIPT_DIR = REVISION_ROOT / "manuscript"
SNAPSHOT_DIR = REVISION_ROOT / "snapshot"

for path in [OUTPUT_TABLES, OUTPUT_FIGURES, REPORTS_DIR, MANUSCRIPT_DIR, SNAPSHOT_DIR]:
    path.mkdir(parents=True, exist_ok=True)

if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.schemas.bottleneck import BottleneckDistribution, BottleneckPrediction
from app.schemas.demo import ScenarioInput
from app.schemas.inventory import InventoryPrediction
from app.schemas.maintenance import MaintenancePrediction
from app.services.bottleneck_service import predict_bottleneck
from app.services.fusion_service import predict_fusion
from app.services.inventory_service import predict_inventory
from app.services.maintenance_service import predict_maintenance
from app.services.scenario_service import (
    compute_scenario,
    get_demo_scenarios,
    to_bottleneck_input,
    to_fusion_input,
    to_inventory_input,
    to_maintenance_input,
)


AIRCRAFT_TYPES = ["A320neo", "A321", "B737-800", "E195-E2"]
ROUTE_CLASSES = ["short_haul", "medium_haul", "long_haul"]
SERVICE_PROFILES = ["standard", "business_priority", "high_turnover", "late_rotation"]
INVENTORY_CATEGORIES = ["beverage", "meal", "duty_free", "wheelchair_kit", "mixed_service"]
COMPONENT_TYPES = [
    "seat_power_unit",
    "galley_chiller",
    "ife_display",
    "lavatory_sensor",
    "cabin_light_driver",
]
REGIMES = [
    "low_risk_routine",
    "maintenance_heavy",
    "inventory_heavy",
    "congestion_delay_heavy",
    "assistance_readiness",
    "composite_multi_risk",
]
ALERT_ORDER = ["Routine", "Watch", "Priority", "Critical"]
BOTTLENECK_ORDER = ["Maintenance", "Catering", "Cleaning", "Boarding", "Refueling"]
ACTION_CATEGORIES = [
    "Maintenance",
    "Inventory/Catering",
    "Assistance",
    "Gate/Turnaround",
    "OCC coordination",
    "Boarding",
    "Refueling",
    "Cleaning",
    "Operator oversight",
]
URGENCY_RANK = {"OK": 0, "PLAN": 1, "SOON": 2, "CRITICAL": 3}
RISK_RANK = {"LOW": 0, "ELEVATED": 1, "HIGH": 2}


def clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def uniform(rng: random.Random, low: float, high: float, digits: int = 2) -> float:
    return round(rng.uniform(low, high), digits)


def randint(rng: random.Random, low: int, high: int) -> int:
    return rng.randint(low, high)


def stress_level(value: float, low_cut: float, high_cut: float) -> str:
    if value < low_cut:
        return "low"
    if value < high_cut:
        return "medium"
    return "high"


def route_for(index: int) -> str:
    routes = [
        "FRA -> BCN",
        "CGN -> PMI",
        "VIE -> MUC",
        "ZRH -> HAM",
        "DUB -> BER",
        "MAD -> FRA",
        "MUC -> LHR",
        "OSL -> BER",
        "AMS -> VIE",
        "CPH -> MUC",
    ]
    return routes[index % len(routes)]


def make_scenario(index: int, regime: str, rng: random.Random) -> ScenarioInput:
    aircraft_type = rng.choice(AIRCRAFT_TYPES)
    route_class = rng.choices(ROUTE_CLASSES, weights=[0.42, 0.38, 0.20], k=1)[0]
    service_profile = rng.choice(SERVICE_PROFILES)
    inventory_category = rng.choice(INVENTORY_CATEGORIES)
    component_type = rng.choice(COMPONENT_TYPES)
    assistance_request = rng.random() < 0.22
    assistance_confirmed = assistance_request and rng.random() < 0.62

    values = {
        "passenger_load": randint(rng, 85, 180),
        "arrival_delay": randint(rng, 0, 18),
        "gate_congestion": uniform(rng, 1.5, 5.8),
        "inventory_shortage_score": uniform(rng, 1.0, 5.5),
        "trolley_availability": uniform(rng, 0.58, 0.95),
        "catering_complexity": uniform(rng, 2.0, 5.8),
        "turnaround_pressure": uniform(rng, 2.0, 5.8),
        "item_criticality": uniform(rng, 2.0, 5.8),
        "malfunction_count": randint(rng, 0, 2),
        "malfunction_severity": uniform(rng, 0.4, 4.8),
        "aircraft_age_cycles": randint(rng, 6000, 22000),
        "flight_duration_hr": uniform(rng, 0.8, 3.2),
        "cabin_temperature": uniform(rng, 20.0, 26.5),
        "humidity": uniform(rng, 22.0, 50.0),
        "cycles_since_install": randint(rng, 300, 2300),
        "wear_index": uniform(rng, 20.0, 62.0),
        "environmental_stress": uniform(rng, 1.2, 5.5),
        "weather_disturbance": uniform(rng, 0.0, 3.5),
        "refueling_required": rng.random() < 0.78,
    }

    if regime == "low_risk_routine":
        values.update(
            passenger_load=randint(rng, 85, 145),
            arrival_delay=randint(rng, 0, 6),
            gate_congestion=uniform(rng, 1.0, 3.8),
            inventory_shortage_score=uniform(rng, 0.3, 2.8),
            trolley_availability=uniform(rng, 0.78, 0.98),
            catering_complexity=uniform(rng, 1.2, 3.8),
            turnaround_pressure=uniform(rng, 1.5, 3.8),
            item_criticality=uniform(rng, 1.5, 4.0),
            malfunction_count=randint(rng, 0, 1),
            malfunction_severity=uniform(rng, 0.1, 2.2),
            cycles_since_install=randint(rng, 120, 1200),
            wear_index=uniform(rng, 8.0, 38.0),
            environmental_stress=uniform(rng, 0.5, 2.8),
            weather_disturbance=uniform(rng, 0.0, 1.2),
        )
        assistance_request = False
        assistance_confirmed = False
    elif regime == "maintenance_heavy":
        values.update(
            malfunction_count=randint(rng, 3, 7),
            malfunction_severity=uniform(rng, 6.2, 9.8),
            cycles_since_install=randint(rng, 2300, 4100),
            wear_index=uniform(rng, 68.0, 97.0),
            environmental_stress=uniform(rng, 5.5, 9.8),
            cabin_temperature=uniform(rng, 25.5, 34.0),
            humidity=uniform(rng, 38.0, 76.0),
            arrival_delay=randint(rng, 6, 28),
            gate_congestion=uniform(rng, 3.8, 7.2),
        )
        component_type = rng.choice(["galley_chiller", "ife_display", "seat_power_unit"])
    elif regime == "inventory_heavy":
        values.update(
            passenger_load=randint(rng, 150, 215),
            inventory_shortage_score=uniform(rng, 6.7, 9.8),
            trolley_availability=uniform(rng, 0.18, 0.48),
            catering_complexity=uniform(rng, 6.0, 9.8),
            turnaround_pressure=uniform(rng, 5.2, 8.8),
            item_criticality=uniform(rng, 6.0, 9.5),
            arrival_delay=randint(rng, 3, 22),
            gate_congestion=uniform(rng, 3.5, 6.8),
            malfunction_count=randint(rng, 0, 2),
            malfunction_severity=uniform(rng, 0.4, 3.5),
        )
        service_profile = rng.choice(["business_priority", "high_turnover", "late_rotation"])
        inventory_category = rng.choice(["beverage", "meal", "mixed_service"])
    elif regime == "congestion_delay_heavy":
        values.update(
            passenger_load=randint(rng, 145, 215),
            arrival_delay=randint(rng, 22, 70),
            gate_congestion=uniform(rng, 6.5, 9.8),
            turnaround_pressure=uniform(rng, 6.4, 9.8),
            weather_disturbance=uniform(rng, 2.0, 8.0),
            inventory_shortage_score=uniform(rng, 2.5, 6.4),
            trolley_availability=uniform(rng, 0.48, 0.85),
            malfunction_count=randint(rng, 0, 2),
            malfunction_severity=uniform(rng, 0.6, 4.2),
        )
    elif regime == "assistance_readiness":
        values.update(
            passenger_load=randint(rng, 90, 180),
            arrival_delay=randint(rng, 0, 18),
            gate_congestion=uniform(rng, 2.2, 6.2),
            inventory_shortage_score=uniform(rng, 0.8, 4.0),
            trolley_availability=uniform(rng, 0.68, 0.96),
            malfunction_count=randint(rng, 0, 1),
            malfunction_severity=uniform(rng, 0.2, 2.8),
        )
        assistance_request = True
        assistance_confirmed = rng.random() < 0.35
        inventory_category = "wheelchair_kit"
    elif regime == "composite_multi_risk":
        values.update(
            passenger_load=randint(rng, 165, 220),
            arrival_delay=randint(rng, 18, 65),
            gate_congestion=uniform(rng, 6.0, 9.8),
            inventory_shortage_score=uniform(rng, 5.8, 9.6),
            trolley_availability=uniform(rng, 0.20, 0.58),
            catering_complexity=uniform(rng, 5.8, 9.8),
            turnaround_pressure=uniform(rng, 6.5, 9.8),
            item_criticality=uniform(rng, 5.8, 9.8),
            malfunction_count=randint(rng, 2, 6),
            malfunction_severity=uniform(rng, 5.2, 9.6),
            cycles_since_install=randint(rng, 1800, 4100),
            wear_index=uniform(rng, 58.0, 96.0),
            environmental_stress=uniform(rng, 4.8, 9.5),
            weather_disturbance=uniform(rng, 1.5, 8.0),
            refueling_required=True,
        )
        assistance_request = rng.random() < 0.55
        assistance_confirmed = assistance_request and rng.random() < 0.55
        service_profile = rng.choice(["business_priority", "high_turnover", "late_rotation"])

    return ScenarioInput(
        id=f"jatm-{index + 1:04d}-{regime.replace('_', '-')}",
        flight_number=f"JATM-{index + 1:04d}",
        aircraft_type=aircraft_type,
        route=route_for(index),
        route_class=route_class,
        eta=f"{14 + (index % 8):02d}:{(index * 7) % 60:02d}",
        gate=f"{chr(65 + index % 4)}{1 + index % 22:02d}",
        passenger_load=int(values["passenger_load"]),
        arrival_delay=int(values["arrival_delay"]),
        gate_congestion=float(values["gate_congestion"]),
        assistance_request=bool(assistance_request),
        assistance_equipment_confirmed=bool(assistance_confirmed),
        inventory_shortage_score=float(values["inventory_shortage_score"]),
        trolley_availability=float(values["trolley_availability"]),
        service_profile=service_profile,
        catering_complexity=float(values["catering_complexity"]),
        turnaround_pressure=float(values["turnaround_pressure"]),
        item_criticality=float(values["item_criticality"]),
        inventory_category=inventory_category,
        malfunction_count=int(values["malfunction_count"]),
        malfunction_severity=float(values["malfunction_severity"]),
        component_type=component_type,
        aircraft_age_cycles=int(values["aircraft_age_cycles"]),
        flight_duration_hr=float(values["flight_duration_hr"]),
        cabin_temperature=float(values["cabin_temperature"]),
        humidity=float(values["humidity"]),
        cycles_since_install=int(values["cycles_since_install"]),
        wear_index=float(values["wear_index"]),
        environmental_stress=float(values["environmental_stress"]),
        weather_disturbance=float(values["weather_disturbance"]),
        refueling_required=bool(values["refueling_required"]),
        cabin_report_notes=[f"Controlled synthetic {regime.replace('_', ' ')} evaluation case."],
        reported_components=[component_type.replace("_", " "), inventory_category.replace("_", " ")],
    )


def generate_scenarios(count: int = 1000, seed: int = 20260608) -> list[tuple[str, ScenarioInput]]:
    rng = random.Random(seed)
    scenarios: list[tuple[str, ScenarioInput]] = []
    for index in range(count):
        regime = REGIMES[index % len(REGIMES)]
        scenarios.append((regime, make_scenario(index, regime, rng)))
    rng.shuffle(scenarios)
    return scenarios


def action_category(action) -> str:
    text = f"{action.team} {action.description}".lower()
    if "maintenance" in text or "technician" in text or "spare component" in text:
        return "Maintenance"
    if "catering" in text or "trolley" in text or "inventory" in text:
        return "Inventory/Catering"
    if "assistance" in text or "wheelchair" in text or "prm" in text:
        return "Assistance"
    if "gate" in text or "turnaround duty" in text:
        return "Gate/Turnaround"
    if "boarding" in text:
        return "Boarding"
    if "fuel" in text or "refuel" in text:
        return "Refueling"
    if "cleaning" in text:
        return "Cleaning"
    if "occ" in text or "cross-team" in text or "supervisor" in text:
        if "maintain operator review" in text:
            return "Operator oversight"
        return "OCC coordination"
    return "Operator oversight"


def action_counts(response) -> Counter:
    return Counter(action_category(action) for action in response.fusion.action_queue)


def top_action_category(response) -> str:
    for action in response.fusion.action_queue:
        if action.priority != "P4":
            return action_category(action)
    return "Operator oversight"


def scenario_input_row(regime: str, scenario: ScenarioInput) -> dict:
    row = scenario.model_dump()
    row["regime"] = regime
    row["arrival_delay_regime"] = stress_level(scenario.arrival_delay, 8, 24)
    row["gate_congestion_regime"] = stress_level(scenario.gate_congestion, 4, 7)
    row["malfunction_count_regime"] = stress_level(scenario.malfunction_count, 2, 4)
    row["malfunction_severity_regime"] = stress_level(scenario.malfunction_severity, 3.5, 6.5)
    row["inventory_shortage_regime"] = stress_level(scenario.inventory_shortage_score, 3.5, 6.5)
    row["trolley_availability_regime"] = "low" if scenario.trolley_availability < 0.45 else "medium" if scenario.trolley_availability < 0.72 else "high"
    row["passenger_load_regime"] = stress_level(scenario.passenger_load, 130, 175)
    row["assistance_readiness_regime"] = (
        "unconfirmed" if scenario.assistance_request and not scenario.assistance_equipment_confirmed else
        "confirmed" if scenario.assistance_request else
        "not_required"
    )
    row["weather_regime"] = stress_level(scenario.weather_disturbance, 2, 5)
    return row


def scenario_output_row(regime: str, scenario: ScenarioInput, response) -> dict:
    probs = response.bottleneck.probabilities.model_dump()
    counts = action_counts(response)
    row = {
        "scenario_id": scenario.id,
        "flight_number": scenario.flight_number,
        "regime": regime,
        "readiness_score": response.fusion.readiness_score,
        "criticality_score": response.fusion.criticality_score,
        "alert_level": response.fusion.alert_level,
        "at_risk": response.fusion.at_risk,
        "assistance_request": scenario.assistance_request,
        "assistance_equipment_confirmed": scenario.assistance_equipment_confirmed,
        "assistance_readiness_state": (
            "unconfirmed" if scenario.assistance_request and not scenario.assistance_equipment_confirmed else
            "confirmed" if scenario.assistance_request else
            "not_required"
        ),
        "action_queue_size": len(response.fusion.action_queue),
        "p1_action_count": sum(1 for action in response.fusion.action_queue if action.priority == "P1"),
        "p2_action_count": sum(1 for action in response.fusion.action_queue if action.priority == "P2"),
        "top_action_category": top_action_category(response),
        "maintenance_failure_probability": response.maintenance.failure_probability,
        "maintenance_urgency": response.maintenance.urgency_class,
        "maintenance_remaining_flights": response.maintenance.remaining_flights_estimate,
        "inventory_shortage_risk": response.inventory.shortage_risk,
        "inventory_risk_level": response.inventory.risk_level,
        "bottleneck_dominant": response.bottleneck.dominant_bottleneck,
        "bottleneck_confidence": response.bottleneck.confidence,
        "expected_delay_risk_min": response.fusion.expected_delay_risk_min,
        "expected_benefit_min": response.fusion.expected_benefit_min,
        "proactive_vs_reactive_delta_min": response.fusion.proactive_vs_reactive_delta_min,
    }
    for label in BOTTLENECK_ORDER:
        row[f"bottleneck_prob_{label.lower()}"] = probs[label]
    for category in ACTION_CATEGORIES:
        row[f"action_count_{category.lower().replace('/', '_').replace(' ', '_')}"] = counts.get(category, 0)
    return row


def run_scenario_sweep(count: int = 1000, seed: int = 20260608) -> tuple[pd.DataFrame, pd.DataFrame]:
    scenarios = generate_scenarios(count=count, seed=seed)
    input_rows = []
    output_rows = []
    for regime, scenario in scenarios:
        response = compute_scenario(scenario)
        input_rows.append(scenario_input_row(regime, scenario))
        output_rows.append(scenario_output_row(regime, scenario, response))
    inputs_df = pd.DataFrame(input_rows).sort_values("id").reset_index(drop=True)
    outputs_df = pd.DataFrame(output_rows).sort_values("scenario_id").reset_index(drop=True)
    inputs_df.to_csv(OUTPUT_TABLES / "scenario_sweep_inputs.csv", index=False)
    outputs_df.to_csv(OUTPUT_TABLES / "scenario_sweep_outputs.csv", index=False)
    write_scenario_summary(outputs_df)
    write_regime_summary(outputs_df)
    plot_scenario_sweep(outputs_df)
    return inputs_df, outputs_df


def write_scenario_summary(outputs_df: pd.DataFrame) -> None:
    rows = [
        {"metric": "scenario_count", "value": len(outputs_df)},
        {"metric": "mean_readiness", "value": round(outputs_df["readiness_score"].mean(), 3)},
        {"metric": "median_readiness", "value": round(outputs_df["readiness_score"].median(), 3)},
        {"metric": "std_readiness", "value": round(outputs_df["readiness_score"].std(), 3)},
        {"metric": "mean_maintenance_failure_probability", "value": round(outputs_df["maintenance_failure_probability"].mean(), 3)},
        {"metric": "mean_inventory_shortage_risk", "value": round(outputs_df["inventory_shortage_risk"].mean(), 3)},
        {"metric": "mean_action_queue_size", "value": round(outputs_df["action_queue_size"].mean(), 3)},
    ]
    for alert in ALERT_ORDER:
        count = int((outputs_df["alert_level"] == alert).sum())
        rows.append({"metric": f"alert_{alert.lower()}_count", "value": count})
        rows.append({"metric": f"alert_{alert.lower()}_share", "value": round(count / len(outputs_df), 4)})
    for label in BOTTLENECK_ORDER:
        rows.append(
            {
                "metric": f"dominant_bottleneck_{label.lower()}_count",
                "value": int((outputs_df["bottleneck_dominant"] == label).sum()),
            }
        )
    for category in ACTION_CATEGORIES:
        col = f"action_count_{category.lower().replace('/', '_').replace(' ', '_')}"
        rows.append({"metric": f"action_category_{category.lower().replace('/', '_').replace(' ', '_')}_count", "value": int(outputs_df[col].sum())})
    pd.DataFrame(rows).to_csv(OUTPUT_TABLES / "scenario_sweep_summary.csv", index=False)


def write_regime_summary(outputs_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for regime, group in outputs_df.groupby("regime"):
        action_totals = {}
        for category in ACTION_CATEGORIES:
            col = f"action_count_{category.lower().replace('/', '_').replace(' ', '_')}"
            action_totals[category] = int(group[col].sum())
        top_actions = ", ".join(f"{name} ({count})" for name, count in Counter(action_totals).most_common(3) if count > 0)
        alert_counts = {alert: int((group["alert_level"] == alert).sum()) for alert in ALERT_ORDER}
        bottleneck_counts = {label: int((group["bottleneck_dominant"] == label).sum()) for label in BOTTLENECK_ORDER}
        confirmed_readiness = group[group["assistance_readiness_state"] == "confirmed"]["readiness_score"]
        unconfirmed_readiness = group[group["assistance_readiness_state"] == "unconfirmed"]["readiness_score"]
        no_request_readiness = group[group["assistance_readiness_state"] == "not_required"]["readiness_score"]
        confirmed_mean = round(confirmed_readiness.mean(), 2) if len(confirmed_readiness) else ""
        unconfirmed_mean = round(unconfirmed_readiness.mean(), 2) if len(unconfirmed_readiness) else ""
        no_request_mean = round(no_request_readiness.mean(), 2) if len(no_request_readiness) else ""
        confirmed_minus_unconfirmed = (
            round(confirmed_readiness.mean() - unconfirmed_readiness.mean(), 2)
            if len(confirmed_readiness) and len(unconfirmed_readiness)
            else ""
        )
        no_request_minus_unconfirmed = (
            round(no_request_readiness.mean() - unconfirmed_readiness.mean(), 2)
            if len(no_request_readiness) and len(unconfirmed_readiness)
            else ""
        )
        row = {
            "regime": regime,
            "scenario_count": len(group),
            "mean_readiness": round(group["readiness_score"].mean(), 2),
            "median_readiness": round(group["readiness_score"].median(), 2),
            "mean_maintenance_failure_probability": round(group["maintenance_failure_probability"].mean(), 3),
            "mean_inventory_shortage_risk": round(group["inventory_shortage_risk"].mean(), 3),
            "mean_action_queue_size": round(group["action_queue_size"].mean(), 2),
            "top_action_categories": top_actions,
            "dominant_bottleneck_mode": group["bottleneck_dominant"].mode().iat[0],
            "assistance_confirmed_mean_readiness": confirmed_mean,
            "assistance_unconfirmed_mean_readiness": unconfirmed_mean,
            "assistance_not_required_mean_readiness": no_request_mean,
            "assistance_confirmed_minus_unconfirmed": confirmed_minus_unconfirmed,
            "assistance_not_required_minus_unconfirmed": no_request_minus_unconfirmed,
        }
        for alert, count in alert_counts.items():
            row[f"alert_{alert.lower()}_count"] = count
            row[f"alert_{alert.lower()}_share"] = round(count / len(group), 3)
        for label, count in bottleneck_counts.items():
            row[f"bottleneck_{label.lower()}_count"] = count
        rows.append(row)
    summary_df = pd.DataFrame(rows).sort_values("regime")
    summary_df.to_csv(OUTPUT_TABLES / "regime_summary_table.csv", index=False)
    write_markdown_table(summary_df, REPORTS_DIR / "regime_summary_table_for_manuscript.md")
    write_docx_table(
        summary_df[
            [
                "regime",
                "scenario_count",
                "mean_readiness",
                "median_readiness",
                "dominant_bottleneck_mode",
                "top_action_categories",
                "mean_maintenance_failure_probability",
                "mean_inventory_shortage_risk",
            ]
        ],
        REPORTS_DIR / "regime_summary_table_for_manuscript.docx",
        "Regime-level summary of the controlled synthetic scenario sweep.",
    )
    plot_regime_figures(outputs_df)
    return summary_df


def write_markdown_table(df: pd.DataFrame, path: Path) -> None:
    headers = [str(col) for col in df.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for _, row in df.iterrows():
        values = [str(row[col]).replace("|", "/") for col in df.columns]
        lines.append("| " + " | ".join(values) + " |")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_docx_table(df: pd.DataFrame, path: Path, title: str) -> None:
    doc = Document()
    set_document_defaults(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True)
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Table Grid"
    for col_index, col_name in enumerate(df.columns):
        cell = table.cell(0, col_index)
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8, bold=True)
    for row_index, (_, row) in enumerate(df.iterrows(), start=1):
        for col_index, col_name in enumerate(df.columns):
            cell = table.cell(row_index, col_index)
            cell.text = str(row[col_name])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=8)
    doc.save(path)


def plot_scenario_sweep(outputs_df: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.hist(outputs_df["readiness_score"], bins=24, color="#2f6f9f", alpha=0.85)
    ax.set_title("Scenario sweep readiness distribution")
    ax.set_xlabel("Fused readiness score")
    ax.set_ylabel("Scenario count")
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "scenario_sweep_readiness_distribution.png", dpi=300)
    plt.close(fig)

    alert_counts = outputs_df["alert_level"].value_counts().reindex(ALERT_ORDER, fill_value=0)
    fig, ax = plt.subplots(figsize=(6.4, 4.0))
    ax.bar(alert_counts.index, alert_counts.values, color=["#4b8f6a", "#d8a33e", "#cf6f33", "#b83c3c"])
    ax.set_title("Scenario sweep alert distribution")
    ax.set_ylabel("Scenario count")
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "scenario_sweep_alert_distribution.png", dpi=300)
    plt.close(fig)

    bottleneck_counts = outputs_df["bottleneck_dominant"].value_counts().reindex(BOTTLENECK_ORDER, fill_value=0)
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    ax.bar(bottleneck_counts.index, bottleneck_counts.values, color="#3f6f7f")
    ax.set_title("Dominant bottleneck distribution")
    ax.set_ylabel("Scenario count")
    ax.tick_params(axis="x", rotation=20)
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "scenario_sweep_bottleneck_distribution.png", dpi=300)
    plt.close(fig)

    heatmap_data = []
    for regime in REGIMES:
        group = outputs_df[outputs_df["regime"] == regime]
        row = []
        for category in ACTION_CATEGORIES:
            col = f"action_count_{category.lower().replace('/', '_').replace(' ', '_')}"
            row.append(float(group[col].mean()) if len(group) else 0.0)
        heatmap_data.append(row)
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    im = ax.imshow(heatmap_data, aspect="auto", cmap="YlGnBu")
    ax.set_yticks(range(len(REGIMES)), [label.replace("_", " ") for label in REGIMES])
    ax.set_xticks(range(len(ACTION_CATEGORIES)), ACTION_CATEGORIES, rotation=35, ha="right")
    ax.set_title("Mean action-category count by regime")
    fig.colorbar(im, ax=ax, label="Mean actions per scenario")
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "scenario_sweep_action_category_heatmap.png", dpi=300)
    plt.close(fig)


def plot_regime_figures(outputs_df: pd.DataFrame) -> None:
    data = [outputs_df[outputs_df["regime"] == regime]["readiness_score"].values for regime in REGIMES]
    fig, ax = plt.subplots(figsize=(8.0, 4.5))
    ax.boxplot(data, labels=[regime.replace("_", "\n") for regime in REGIMES], patch_artist=True)
    ax.set_title("Readiness by synthetic operational regime")
    ax.set_ylabel("Fused readiness score")
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "regime_readiness_boxplot.png", dpi=300)
    plt.close(fig)

    stacked = []
    for regime in REGIMES:
        group = outputs_df[outputs_df["regime"] == regime]
        total = max(len(group), 1)
        stacked.append([(group["alert_level"] == alert).sum() / total for alert in ALERT_ORDER])
    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    bottom = [0.0] * len(REGIMES)
    colors = ["#4b8f6a", "#d8a33e", "#cf6f33", "#b83c3c"]
    x_labels = [regime.replace("_", "\n") for regime in REGIMES]
    for alert, color, values in zip(ALERT_ORDER, colors, zip(*stacked)):
        values = list(values)
        ax.bar(x_labels, values, bottom=bottom, label=alert, color=color)
        bottom = [b + v for b, v in zip(bottom, values)]
    ax.set_title("Alert distribution by regime")
    ax.set_ylabel("Share of scenarios")
    ax.legend(ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.13))
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "regime_alert_stacked_bar.png", dpi=300)
    plt.close(fig)


def is_non_decreasing(values: list[float], tolerance: float = 1e-9) -> bool:
    return all(values[i] + tolerance >= values[i - 1] for i in range(1, len(values)))


def is_non_increasing(values: list[float], tolerance: float = 1e-9) -> bool:
    return all(values[i] <= values[i - 1] + tolerance for i in range(1, len(values)))


def base_sanity_scenario() -> ScenarioInput:
    return ScenarioInput(
        id="sanity-base",
        flight_number="SANITY-BASE",
        aircraft_type="A320neo",
        route="FRA -> BCN",
        route_class="medium_haul",
        eta="16:00",
        gate="A10",
        passenger_load=150,
        arrival_delay=10,
        gate_congestion=4.8,
        assistance_request=False,
        assistance_equipment_confirmed=False,
        inventory_shortage_score=3.2,
        trolley_availability=0.74,
        service_profile="standard",
        catering_complexity=4.2,
        turnaround_pressure=4.4,
        item_criticality=4.0,
        inventory_category="mixed_service",
        malfunction_count=1,
        malfunction_severity=2.8,
        component_type="seat_power_unit",
        aircraft_age_cycles=15000,
        flight_duration_hr=1.8,
        cabin_temperature=23.5,
        humidity=36,
        cycles_since_install=1200,
        wear_index=45,
        environmental_stress=3.2,
        weather_disturbance=1.0,
        refueling_required=True,
    )


def scenario_copy(scenario: ScenarioInput, **updates) -> ScenarioInput:
    return scenario.model_copy(update=updates)


def run_monotonicity_checks() -> tuple[pd.DataFrame, pd.DataFrame]:
    base = base_sanity_scenario()
    detail_rows = []
    summary_rows = []

    tests = [
        ("malfunction_severity", [0, 2, 4, 6, 8, 10], "maintenance_and_readiness"),
        ("malfunction_count", [0, 1, 2, 3, 5, 7], "readiness_non_increasing"),
        ("inventory_shortage_score", [0, 2, 4, 6, 8, 10], "inventory_risk_non_decreasing"),
        ("trolley_availability", [1.0, 0.8, 0.6, 0.4, 0.25], "inventory_risk_non_decreasing_when_availability_decreases"),
        ("arrival_delay", [0, 5, 15, 30, 50, 80], "readiness_non_increasing"),
        ("gate_congestion", [0.5, 2, 4, 6, 8, 10], "readiness_non_increasing"),
    ]

    for variable, values, expectation in tests:
        readiness_values = []
        failure_values = []
        urgency_values = []
        inventory_values = []
        inventory_rank_values = []
        for value in values:
            scenario = scenario_copy(base, **{variable: value})
            if variable == "trolley_availability":
                scenario = scenario_copy(base, trolley_availability=value)
            response = compute_scenario(scenario)
            readiness_values.append(response.fusion.readiness_score)
            failure_values.append(response.maintenance.failure_probability)
            urgency_values.append(URGENCY_RANK[response.maintenance.urgency_class])
            inventory_values.append(response.inventory.shortage_risk)
            inventory_rank_values.append(RISK_RANK[response.inventory.risk_level])
            detail_rows.append(
                {
                    "test": variable,
                    "input_value": value,
                    "readiness_score": response.fusion.readiness_score,
                    "maintenance_failure_probability": response.maintenance.failure_probability,
                    "maintenance_urgency": response.maintenance.urgency_class,
                    "maintenance_urgency_rank": URGENCY_RANK[response.maintenance.urgency_class],
                    "inventory_shortage_risk": response.inventory.shortage_risk,
                    "inventory_risk_level": response.inventory.risk_level,
                    "inventory_risk_rank": RISK_RANK[response.inventory.risk_level],
                    "alert_level": response.fusion.alert_level,
                }
            )
        if variable == "malfunction_severity":
            passed = is_non_decreasing(failure_values) and is_non_decreasing(urgency_values) and is_non_increasing(readiness_values)
            criterion = "failure risk and urgency non-decreasing; readiness non-increasing"
        elif variable == "malfunction_count":
            passed = is_non_increasing(readiness_values)
            criterion = "readiness non-increasing"
        elif variable == "inventory_shortage_score":
            passed = is_non_decreasing(inventory_values) and is_non_decreasing(inventory_rank_values)
            criterion = "inventory risk non-decreasing"
        elif variable == "trolley_availability":
            passed = is_non_decreasing(inventory_values) and is_non_decreasing(inventory_rank_values)
            criterion = "inventory risk non-decreasing as availability decreases"
        else:
            passed = is_non_increasing(readiness_values)
            criterion = "readiness non-increasing"
        summary_rows.append(
            {
                "test": variable,
                "criterion": criterion,
                "passed": bool(passed),
                "min_readiness": min(readiness_values),
                "max_readiness": max(readiness_values),
                "notes": "Controlled one-factor sweep; not field validation.",
            }
        )

    no_request = compute_scenario(scenario_copy(base, assistance_request=False, assistance_equipment_confirmed=False))
    confirmed = compute_scenario(scenario_copy(base, assistance_request=True, assistance_equipment_confirmed=True))
    unconfirmed = compute_scenario(scenario_copy(base, assistance_request=True, assistance_equipment_confirmed=False))
    assistance_passed = unconfirmed.fusion.readiness_score <= min(no_request.fusion.readiness_score, confirmed.fusion.readiness_score)
    for label, response in [
        ("no_request", no_request),
        ("confirmed", confirmed),
        ("unconfirmed", unconfirmed),
    ]:
        detail_rows.append(
            {
                "test": "assistance_request",
                "input_value": label,
                "readiness_score": response.fusion.readiness_score,
                "maintenance_failure_probability": response.maintenance.failure_probability,
                "maintenance_urgency": response.maintenance.urgency_class,
                "maintenance_urgency_rank": URGENCY_RANK[response.maintenance.urgency_class],
                "inventory_shortage_risk": response.inventory.shortage_risk,
                "inventory_risk_level": response.inventory.risk_level,
                "inventory_risk_rank": RISK_RANK[response.inventory.risk_level],
                "alert_level": response.fusion.alert_level,
            }
        )
    summary_rows.append(
        {
            "test": "assistance_request",
            "criterion": "unconfirmed assistance request does not improve readiness relative to no request or confirmed support",
            "passed": bool(assistance_passed),
            "min_readiness": min(no_request.fusion.readiness_score, confirmed.fusion.readiness_score, unconfirmed.fusion.readiness_score),
            "max_readiness": max(no_request.fusion.readiness_score, confirmed.fusion.readiness_score, unconfirmed.fusion.readiness_score),
            "notes": "Rule-based assistance-readiness check.",
        }
    )

    low_case = compute_scenario(make_scenario(9001, "low_risk_routine", random.Random(42)))
    composite_case = compute_scenario(make_scenario(9002, "composite_multi_risk", random.Random(43)))
    composite_passed = composite_case.fusion.readiness_score < low_case.fusion.readiness_score
    for label, response in [("low_risk_routine", low_case), ("composite_multi_risk", composite_case)]:
        detail_rows.append(
            {
                "test": "composite_vs_low_risk",
                "input_value": label,
                "readiness_score": response.fusion.readiness_score,
                "maintenance_failure_probability": response.maintenance.failure_probability,
                "maintenance_urgency": response.maintenance.urgency_class,
                "maintenance_urgency_rank": URGENCY_RANK[response.maintenance.urgency_class],
                "inventory_shortage_risk": response.inventory.shortage_risk,
                "inventory_risk_level": response.inventory.risk_level,
                "inventory_risk_rank": RISK_RANK[response.inventory.risk_level],
                "alert_level": response.fusion.alert_level,
            }
        )
    summary_rows.append(
        {
            "test": "composite_vs_low_risk",
            "criterion": "composite multi-risk case has lower readiness than isolated low-risk case",
            "passed": bool(composite_passed),
            "min_readiness": min(low_case.fusion.readiness_score, composite_case.fusion.readiness_score),
            "max_readiness": max(low_case.fusion.readiness_score, composite_case.fusion.readiness_score),
            "notes": "Scenario-class comparison; not empirical validation.",
        }
    )

    detail_df = pd.DataFrame(detail_rows)
    summary_df = pd.DataFrame(summary_rows)
    detail_df.to_csv(OUTPUT_TABLES / "monotonicity_test_results.csv", index=False)
    summary_df.to_csv(OUTPUT_TABLES / "monotonicity_summary.csv", index=False)
    write_monotonicity_summary(summary_df)
    plot_sensitivity_curves(detail_df)
    return detail_df, summary_df


def write_monotonicity_summary(summary_df: pd.DataFrame) -> None:
    lines = [
        "# Monotonicity and Sanity-Check Summary",
        "",
        "These checks evaluate controlled system behavior under synthetic one-factor sweeps. They are not real-world validation.",
        "",
    ]
    for _, row in summary_df.iterrows():
        status = "PASS" if row["passed"] else "FAIL"
        lines.append(f"- {status}: {row['test']} - {row['criterion']}.")
    failures = summary_df[summary_df["passed"] == False]  # noqa: E712
    if len(failures):
        lines.extend(
            [
                "",
                "Failures should be reported rather than hidden. A failure means the current trained-module plus fusion behavior is not strictly monotonic under that synthetic one-factor sweep.",
            ]
        )
    (REPORTS_DIR / "monotonicity_summary.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def plot_sensitivity_curves(detail_df: pd.DataFrame) -> None:
    variables = [
        "malfunction_severity",
        "malfunction_count",
        "inventory_shortage_score",
        "trolley_availability",
        "arrival_delay",
        "gate_congestion",
    ]
    fig, axes = plt.subplots(2, 3, figsize=(12, 6.5))
    for ax, variable in zip(axes.flat, variables):
        group = detail_df[detail_df["test"] == variable].copy()
        ax.plot(group["input_value"], group["readiness_score"], marker="o", label="Readiness")
        if variable in {"malfunction_severity", "malfunction_count"}:
            ax2 = ax.twinx()
            ax2.plot(group["input_value"], group["maintenance_failure_probability"], marker="s", color="#b83c3c", label="Failure risk")
            ax2.set_ylim(0, 1)
        elif variable in {"inventory_shortage_score", "trolley_availability"}:
            ax2 = ax.twinx()
            ax2.plot(group["input_value"], group["inventory_shortage_risk"], marker="s", color="#b83c3c", label="Inventory risk")
            ax2.set_ylim(0, 1)
        ax.set_title(variable.replace("_", " "))
        ax.set_ylabel("Readiness")
        ax.grid(alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "sensitivity_curves.png", dpi=300)
    plt.close(fig)

    for variable in variables:
        group = detail_df[detail_df["test"] == variable].copy()
        fig, ax = plt.subplots(figsize=(6.2, 3.8))
        ax.plot(group["input_value"], group["readiness_score"], marker="o", color="#2f6f9f")
        ax.set_title(f"Sensitivity: {variable.replace('_', ' ')}")
        ax.set_xlabel(variable.replace("_", " "))
        ax.set_ylabel("Readiness score")
        ax.grid(alpha=0.22)
        fig.tight_layout()
        fig.savefig(OUTPUT_FIGURES / f"sensitivity_{variable}.png", dpi=300)
        plt.close(fig)


def neutral_maintenance(original: MaintenancePrediction) -> MaintenancePrediction:
    return original.model_copy(
        update={
            "failure_probability": 0.05,
            "confidence": max(original.confidence, 0.75),
            "remaining_flights_estimate": max(original.remaining_flights_estimate, 48),
            "urgency_class": "OK",
            "anomaly_flag": False,
            "top_factors": ["Maintenance influence neutralized for ablation."],
            "explanation": "Maintenance influence neutralized for ablation.",
            "recommended_action": "No maintenance action generated in this ablation.",
        }
    )


def neutral_inventory(original: InventoryPrediction) -> InventoryPrediction:
    return original.model_copy(
        update={
            "shortage_risk": 0.05,
            "risk_level": "LOW",
            "trolley_status": "Available",
            "top_factors": ["Inventory influence neutralized for ablation."],
            "recommendation": "No inventory action generated in this ablation.",
            "explanation": "Inventory influence neutralized for ablation.",
        }
    )


def neutral_bottleneck(original: BottleneckPrediction) -> BottleneckPrediction:
    distribution = BottleneckDistribution(Maintenance=0.2, Catering=0.2, Cleaning=0.2, Boarding=0.2, Refueling=0.2)
    return original.model_copy(
        update={
            "probabilities": distribution,
            "dominant_bottleneck": "Cleaning",
            "confidence": 0.2,
            "top_contributing_factors": ["Bottleneck influence neutralized for ablation."],
            "explanation": "Bottleneck influence neutralized for ablation.",
            "recommended_response": "No bottleneck-specific response generated in this ablation.",
        }
    )


def fused_for_configuration(scenario: ScenarioInput, config: str):
    response = compute_scenario(scenario)
    scenario_for_fusion = scenario
    maintenance = response.maintenance
    inventory = response.inventory
    bottleneck = response.bottleneck
    if config == "full_system":
        return response
    if config == "no_maintenance_in_fusion":
        maintenance = neutral_maintenance(maintenance)
    elif config == "no_inventory_in_fusion":
        inventory = neutral_inventory(inventory)
    elif config == "no_bottleneck_in_fusion":
        bottleneck = neutral_bottleneck(bottleneck)
    elif config == "no_assistance_in_fusion":
        scenario_for_fusion = scenario.model_copy(update={"assistance_request": False, "assistance_equipment_confirmed": False})
    else:
        raise ValueError(f"Unknown ablation config: {config}")
    fusion = predict_fusion(
        to_fusion_input(
            scenario_for_fusion,
            maintenance=maintenance,
            inventory=inventory,
            bottleneck=bottleneck,
        )
    )
    return response.model_copy(update={"scenario": scenario_for_fusion, "maintenance": maintenance, "inventory": inventory, "bottleneck": bottleneck, "fusion": fusion})


def run_ablation_baseline(seed: int = 20260608) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    inputs_path = OUTPUT_TABLES / "scenario_sweep_inputs.csv"
    if inputs_path.exists():
        scenarios = load_scenarios_from_inputs(inputs_path)
    else:
        scenarios = generate_scenarios(1000, seed)
    configs = [
        "full_system",
        "no_maintenance_in_fusion",
        "no_inventory_in_fusion",
        "no_bottleneck_in_fusion",
        "no_assistance_in_fusion",
    ]
    detail_rows = []
    full_lookup = {}
    for regime, scenario in scenarios:
        full = fused_for_configuration(scenario, "full_system")
        full_lookup[scenario.id] = full
        for config in configs:
            response = full if config == "full_system" else fused_for_configuration(scenario, config)
            detail_rows.append(
                {
                    "scenario_id": scenario.id,
                    "regime": regime,
                    "configuration": config,
                    "readiness_score": response.fusion.readiness_score,
                    "readiness_delta_from_full": round(response.fusion.readiness_score - full.fusion.readiness_score, 3),
                    "alert_level": response.fusion.alert_level,
                    "action_queue_size": len(response.fusion.action_queue),
                    "critical_case": response.fusion.alert_level == "Critical",
                    "priority_case": response.fusion.alert_level == "Priority",
                    "top_action_category": top_action_category(response),
                    "full_top_action_category": top_action_category(full),
                    "top_action_changed": top_action_category(response) != top_action_category(full),
                }
            )
    detail_df = pd.DataFrame(detail_rows)
    summary_rows = []
    for config, group in detail_df.groupby("configuration"):
        row = {
            "configuration": config,
            "scenario_count": len(group),
            "mean_readiness": round(group["readiness_score"].mean(), 3),
            "mean_readiness_delta_from_full": round(group["readiness_delta_from_full"].mean(), 3),
            "critical_cases": int(group["critical_case"].sum()),
            "priority_cases": int(group["priority_case"].sum()),
            "mean_action_queue_size": round(group["action_queue_size"].mean(), 3),
            "top_action_changed_cases": int(group["top_action_changed"].sum()),
        }
        for alert in ALERT_ORDER:
            row[f"alert_{alert.lower()}_count"] = int((group["alert_level"] == alert).sum())
        for category in ACTION_CATEGORIES:
            row[f"top_action_{category.lower().replace('/', '_').replace(' ', '_')}_count"] = int((group["top_action_category"] == category).sum())
        summary_rows.append(row)
    summary_df = pd.DataFrame(summary_rows).sort_values("configuration")
    detail_df.to_csv(OUTPUT_TABLES / "ablation_action_shift_table.csv", index=False)
    summary_df.to_csv(OUTPUT_TABLES / "ablation_summary.csv", index=False)
    baseline_df = run_baseline_comparison(scenarios)
    plot_ablation_figures(summary_df, detail_df)
    plot_baseline_figure(baseline_df)
    return summary_df, detail_df, baseline_df


def load_scenarios_from_inputs(path: Path) -> list[tuple[str, ScenarioInput]]:
    df = pd.read_csv(path)
    scenarios = []
    model_fields = set(ScenarioInput.model_fields.keys())
    for _, row in df.iterrows():
        payload = {key: row[key] for key in model_fields if key in row}
        for bool_key in ["assistance_request", "assistance_equipment_confirmed", "refueling_required"]:
            payload[bool_key] = bool(payload[bool_key])
        for list_key in ["cabin_report_notes", "reported_components"]:
            value = payload.get(list_key, [])
            if isinstance(value, str):
                payload[list_key] = [value] if value else []
        scenarios.append((row["regime"], ScenarioInput(**payload)))
    return scenarios


def alert_from_readiness(readiness: float) -> str:
    if readiness <= 38:
        return "Critical"
    if readiness <= 56:
        return "Priority"
    if readiness <= 74:
        return "Watch"
    return "Routine"


def baseline_prediction(scenario: ScenarioInput, baseline: str) -> dict:
    response = compute_scenario(scenario)
    if baseline == "full_integrated_occ_fusion":
        return {
            "readiness_score": response.fusion.readiness_score,
            "alert_level": response.fusion.alert_level,
            "dominant_signal": response.bottleneck.dominant_bottleneck,
            "action_category": top_action_category(response),
        }
    if baseline == "static_delay_gate":
        readiness = 100 - min(24, scenario.arrival_delay * 0.42) - scenario.gate_congestion * 3.2
        readiness = round(clamp(readiness, 5, 99), 1)
        return {
            "readiness_score": readiness,
            "alert_level": alert_from_readiness(readiness),
            "dominant_signal": "Delay/Gate",
            "action_category": "Gate/Turnaround",
        }
    if baseline == "maintenance_only":
        m = response.maintenance
        readiness = 100 - {"OK": 4, "PLAN": 15, "SOON": 28, "CRITICAL": 45}[m.urgency_class] - m.failure_probability * 18
        readiness = round(clamp(readiness, 5, 99), 1)
        return {
            "readiness_score": readiness,
            "alert_level": alert_from_readiness(readiness),
            "dominant_signal": m.urgency_class,
            "action_category": "Maintenance" if m.urgency_class != "OK" else "Operator oversight",
        }
    if baseline == "inventory_only":
        inv = response.inventory
        readiness = 100 - inv.shortage_risk * 55
        readiness = round(clamp(readiness, 5, 99), 1)
        return {
            "readiness_score": readiness,
            "alert_level": alert_from_readiness(readiness),
            "dominant_signal": inv.risk_level,
            "action_category": "Inventory/Catering" if inv.risk_level != "LOW" else "Operator oversight",
        }
    if baseline == "bottleneck_only":
        bottleneck_max = max(response.bottleneck.probabilities.model_dump().values())
        readiness = 100 - bottleneck_max * 62
        readiness = round(clamp(readiness, 5, 99), 1)
        return {
            "readiness_score": readiness,
            "alert_level": alert_from_readiness(readiness),
            "dominant_signal": response.bottleneck.dominant_bottleneck,
            "action_category": response.bottleneck.dominant_bottleneck,
        }
    raise ValueError(baseline)


def run_baseline_comparison(scenarios: list[tuple[str, ScenarioInput]]) -> pd.DataFrame:
    baselines = [
        "static_delay_gate",
        "maintenance_only",
        "inventory_only",
        "bottleneck_only",
        "full_integrated_occ_fusion",
    ]
    rows = []
    detail_rows = []
    for baseline in baselines:
        predictions = []
        for regime, scenario in scenarios:
            prediction = baseline_prediction(scenario, baseline)
            prediction.update({"scenario_id": scenario.id, "regime": regime, "baseline": baseline})
            predictions.append(prediction)
            detail_rows.append(prediction)
        frame = pd.DataFrame(predictions)
        row = {
            "baseline": baseline,
            "scenario_count": len(frame),
            "mean_readiness": round(frame["readiness_score"].mean(), 3),
            "std_readiness": round(frame["readiness_score"].std(), 3),
            "unique_alert_levels": int(frame["alert_level"].nunique()),
            "unique_dominant_signals": int(frame["dominant_signal"].nunique()),
            "unique_action_categories": int(frame["action_category"].nunique()),
            "critical_cases": int((frame["alert_level"] == "Critical").sum()),
            "priority_cases": int((frame["alert_level"] == "Priority").sum()),
        }
        for alert in ALERT_ORDER:
            row[f"alert_{alert.lower()}_count"] = int((frame["alert_level"] == alert).sum())
        rows.append(row)
    pd.DataFrame(detail_rows).to_csv(OUTPUT_TABLES / "baseline_comparison_detail.csv", index=False)
    baseline_df = pd.DataFrame(rows)
    baseline_df.to_csv(OUTPUT_TABLES / "baseline_comparison.csv", index=False)
    note = (
        "# Baseline Comparison Interpretation\n\n"
        "The baseline comparison is a controlled synthetic behavior comparison. It does not prove real operational superiority. "
        "The useful question is whether the full integrated OCC fusion creates richer cross-domain differentiation than single-signal baselines.\n"
    )
    (REPORTS_DIR / "baseline_interpretation_note.md").write_text(note, encoding="utf-8")
    return baseline_df


def plot_ablation_figures(summary_df: pd.DataFrame, detail_df: pd.DataFrame) -> None:
    configs = [config for config in summary_df["configuration"] if config != "full_system"]
    values = [
        float(summary_df.loc[summary_df["configuration"] == config, "mean_readiness_delta_from_full"].iat[0])
        for config in configs
    ]
    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.bar([c.replace("_", "\n") for c in configs], values, color="#3f6f7f")
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Ablation mean readiness change relative to full system")
    ax.set_ylabel("Readiness delta")
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "ablation_readiness_delta.png", dpi=300)
    plt.close(fig)

    pivot = detail_df.pivot_table(index="configuration", columns="alert_level", values="scenario_id", aggfunc="count", fill_value=0)
    pivot = pivot.reindex(columns=ALERT_ORDER, fill_value=0)
    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    bottom = [0] * len(pivot)
    colors = ["#4b8f6a", "#d8a33e", "#cf6f33", "#b83c3c"]
    x = [idx.replace("_", "\n") for idx in pivot.index]
    for alert, color in zip(ALERT_ORDER, colors):
        values = list(pivot[alert])
        ax.bar(x, values, bottom=bottom, label=alert, color=color)
        bottom = [b + v for b, v in zip(bottom, values)]
    ax.set_title("Ablation alert distribution")
    ax.set_ylabel("Scenario count")
    ax.legend(ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.15))
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "ablation_alert_distribution.png", dpi=300)
    plt.close(fig)


def plot_baseline_figure(baseline_df: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(8.0, 4.3))
    ax.bar(
        [item.replace("_", "\n") for item in baseline_df["baseline"]],
        baseline_df["unique_action_categories"],
        color="#2f6f9f",
    )
    ax.set_title("Cross-domain differentiation by baseline")
    ax.set_ylabel("Unique action categories")
    ax.grid(axis="y", alpha=0.22)
    fig.tight_layout()
    fig.savefig(OUTPUT_FIGURES / "baseline_comparison_figure.png", dpi=300)
    plt.close(fig)


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        style_obj = document.styles[style_name]
        style_obj.font.name = "Times New Roman"
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def set_run_font(run, size: int = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_numbers(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def file_sha256(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def iter_snapshot_files() -> Iterable[Path]:
    skip_dirs = {
        ".venv",
        "node_modules",
        "__pycache__",
        "dustbin",
        "github_release",
        "_karpathy_weekly_inventory",
        "paper_revision_jatm",
    }
    for path in REPO_ROOT.rglob("*"):
        if not path.is_file():
            continue
        parts = set(path.relative_to(REPO_ROOT).parts)
        if parts & skip_dirs:
            continue
        yield path


def write_snapshot_manifest(path: Path) -> pd.DataFrame:
    rows = []
    for file_path in sorted(iter_snapshot_files(), key=lambda item: str(item).lower()):
        rel = str(file_path.relative_to(REPO_ROOT))
        rows.append(
            {
                "relative_path": rel,
                "size_bytes": file_path.stat().st_size,
                "sha256": file_sha256(file_path),
            }
        )
    df = pd.DataFrame(rows)
    df.to_csv(path, index=False)
    return df


def compare_snapshot(baseline_path: Path, current_path: Path, report_path: Path) -> pd.DataFrame:
    baseline = pd.read_csv(baseline_path)
    current = write_snapshot_manifest(current_path)
    merged = baseline.merge(current, on="relative_path", how="outer", suffixes=("_baseline", "_current"), indicator=True)
    rows = []
    for _, row in merged.iterrows():
        status = "unchanged"
        if row["_merge"] == "left_only":
            status = "deleted"
        elif row["_merge"] == "right_only":
            status = "added"
        elif row["sha256_baseline"] != row["sha256_current"]:
            status = "modified"
        if status != "unchanged":
            rows.append(
                {
                    "relative_path": row["relative_path"],
                    "status": status,
                    "sha256_baseline": row.get("sha256_baseline", ""),
                    "sha256_current": row.get("sha256_current", ""),
                }
            )
    result = pd.DataFrame(rows, columns=["relative_path", "status", "sha256_baseline", "sha256_current"])
    result.to_csv(report_path, index=False)
    return result


def load_metric(summary_path: Path, metric: str, default: str = "NA") -> str:
    if not summary_path.exists():
        return default
    df = pd.read_csv(summary_path)
    match = df[df["metric"] == metric]
    if len(match):
        return str(match["value"].iat[0])
    return default


def read_csv_safe(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)
